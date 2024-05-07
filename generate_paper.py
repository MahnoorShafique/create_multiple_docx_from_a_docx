import csv
import os
from docx import Document
from docx.shared import Pt

def write_docx_file(file_path, document):
    document.save(file_path)

def replace_text_in_runs(paragraph, replacements):
    # Iterate over all runs in a paragraph
    for run in paragraph.runs:
        original_text = run.text
        for key, value in replacements.items():
            if key in original_text:
                # Replace text safely
                run.text = original_text.replace(key, value)

def process_all_paragraphs(doc, replacements):
    # Replace text in main document body paragraphs
    for paragraph in doc.paragraphs:
        replace_text_in_runs(paragraph, replacements)

    # Replace text in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_text_in_runs(paragraph, replacements)

    # Replace text in headers and footers
    for section in doc.sections:
        for header in section.header.paragraphs:
            replace_text_in_runs(header, replacements)
        for footer in section.footer.paragraphs:
            replace_text_in_runs(footer, replacements)

english_exam_file_path = "Checkpoint- English Grade 1.docx"
csv_file_path = "data-1715068243851.csv"
output_directory = "english_paper_dir"

if not os.path.exists(output_directory):
    os.makedirs(output_directory)

students_by_emis_class = {}

with open(csv_file_path, 'r', encoding='utf-8') as csv_file:
    csv_reader = csv.DictReader(csv_file)
    for row in csv_reader:
        emis = row['emis']
        class_id = row['class_id']
        if (emis, class_id) not in students_by_emis_class:
            students_by_emis_class[(emis, class_id)] = []
        students_by_emis_class[(emis, class_id)].append(row)

for (emis, class_id), students in students_by_emis_class.items():
    emis_class_folder = os.path.join(output_directory, f"EMIS_{emis}", f"Class_{class_id}")
    if not os.path.exists(emis_class_folder):
        os.makedirs(emis_class_folder)

    for student in students:
        student_name = student['student_name']
        admission_number = student['admission_number']
        output_file_name = f"english_{admission_number}.docx"
        output_file_path = os.path.join(emis_class_folder, output_file_name)
        doc = Document(english_exam_file_path)

        replacements = {
            'student_name': student_name,
            'admission_number': admission_number,
            'emis': emis
        }

        process_all_paragraphs(doc, replacements)
        write_docx_file(output_file_path, doc)

print("Doc files generated successfully.")
